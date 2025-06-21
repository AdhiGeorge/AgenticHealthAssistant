import streamlit as st
from medical_analysis.agents.chat_manager import ChatManager
from medical_analysis.utils.document_extractor import extract_text_from_document
from medical_analysis.utils.logger import get_logger
from medical_analysis.utils.db import get_conversations_for_session, get_conversation_context
import time
import uuid
from streamlit_lottie import st_lottie
import requests
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# --- LOTTIE ANIMATION LOADER ---
def load_lottieurl(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

# Lottie animation URLs
LOTTIE_MEDICAL = "https://assets2.lottiefiles.com/packages/lf20_ktwnwv5m.json"
LOTTIE_CHAT = "https://assets2.lottiefiles.com/packages/lf20_4kx2q32n.json"
LOTTIE_LOADING = "https://assets2.lottiefiles.com/packages/lf20_usmfx6bp.json"

st.set_page_config(
    page_title="🏥 Medical Diagnostics AI Chatbot",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="expanded"
)

logger = get_logger("streamlit_app")
chat_manager = ChatManager()

# --- SESSION STATE ---
if 'session_id' not in st.session_state:
    st.session_state['session_id'] = str(uuid.uuid4())
if 'conversation_id' not in st.session_state:
    st.session_state['conversation_id'] = None
if 'chat_history' not in st.session_state:
    st.session_state['chat_history'] = []
if 'analysis_report' not in st.session_state:
    st.session_state['analysis_report'] = None
if 'original_data' not in st.session_state:
    st.session_state['original_data'] = None
if 'conversations' not in st.session_state:
    st.session_state['conversations'] = []

# --- SIDEBAR: Conversation List ---
st.sidebar.markdown("<h2 style='margin-bottom:0.5em'>🗂️ Conversations</h2>", unsafe_allow_html=True)
convs = get_conversations_for_session(st.session_state['session_id'])
conv_titles = [c[1] for c in convs]
conv_ids = [c[0] for c in convs]
selected_idx = 0
if st.session_state['conversation_id'] in conv_ids:
    selected_idx = conv_ids.index(st.session_state['conversation_id'])
selected = st.sidebar.radio(
    label="Select a conversation:",
    options=range(len(conv_titles)),
    format_func=lambda i: conv_titles[i],
    index=selected_idx if conv_titles else 0,
    key="sidebar_conversation_radio"
) if conv_titles else None
if conv_titles:
    st.session_state['conversation_id'] = conv_ids[selected]
    # Load context for selected conversation
    context = get_conversation_context(st.session_state['conversation_id'])
    if context:
        st.session_state['analysis_report'] = context.get('analysis')
        st.session_state['original_data'] = context.get('original')
        st.session_state['chat_history'] = [
            {'role': msg[1], 'content': msg[2]} for msg in context.get('chat_history', [])
        ]
else:
    context = None

if st.sidebar.button("➕ New Conversation"):
    st.session_state['conversation_id'] = chat_manager.start_conversation(st.session_state['session_id'])
    st.session_state['analysis_report'] = None
    st.session_state['original_data'] = None
    st.session_state['chat_history'] = []
    st.rerun()

st.sidebar.markdown("<hr style='margin:1em 0'>", unsafe_allow_html=True)
st.sidebar.markdown("<small>All data is processed locally and securely.</small>", unsafe_allow_html=True)

# --- MAIN AREA: Claude-style single window ---
st.markdown(
    """
    <style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1976d2;
        text-align: center;
        margin-bottom: 0.5em;
        letter-spacing: 1px;
        text-shadow: 0 2px 8px #e3f2fd;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #555;
        text-align: center;
        margin-bottom: 1.5em;
    }
    .chat-bubble-user {
        background: linear-gradient(90deg, #e3f2fd 0%, #bbdefb 100%);
        border-radius: 1.2em 1.2em 0 1.2em;
        padding: 1em;
        margin: 0.5em 0 0.5em 2em;
        max-width: 70%;
        align-self: flex-end;
        box-shadow: 0 2px 8px #e3f2fd;
    }
    .chat-bubble-ai {
        background: linear-gradient(90deg, #fffde7 0%, #fff9c4 100%);
        border-radius: 1.2em 1.2em 1.2em 0;
        padding: 1em;
        margin: 0.5em 2em 0.5em 0;
        max-width: 70%;
        align-self: flex-start;
        box-shadow: 0 2px 8px #fffde7;
    }
    .chat-container {
        display: flex;
        flex-direction: column;
        gap: 0.5em;
        min-height: 400px;
        margin-bottom: 1em;
    }
    .stButton>button {
        background: #1976d2;
        color: white;
        font-weight: bold;
        border-radius: 0.5em;
        padding: 0.5em 1.5em;
        border: none;
        transition: background 0.3s;
    }
    .stButton>button:hover {
        background: #1565c0;
    }
    </style>
    <div class="main-header">🏥 Medical Diagnostics AI Chatbot</div>
    <div class="sub-header">Advanced multi-agent medical analysis, chat, and Q&A. Powered by Gemini, Groq, and Qdrant.</div>
    """,
    unsafe_allow_html=True
)

# --- Upload/Input if no analysis yet ---
if not st.session_state['analysis_report']:
    st_lottie(load_lottieurl(LOTTIE_MEDICAL), height=180, key="main-medical-anim")
    st.subheader("Upload a Medical Document or Enter Text")
    uploaded_file = st.file_uploader("Choose a file (PDF, DOCX, TXT, CSV, XLSX, JPG, PNG)", type=["pdf", "docx", "txt", "csv", "xlsx", "jpg", "jpeg", "png"])
    text_input = st.text_area("Or paste/type your medical text here:", height=120)
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Analyze Document") and uploaded_file is not None:
            with st.spinner("Extracting text from document..."):
                try:
                    temp_path = f"/tmp/{uploaded_file.name}"
                    with open(temp_path, "wb") as f:
                        f.write(uploaded_file.getbuffer())
                    text = extract_text_from_document(temp_path)
                    st.session_state['original_data'] = text
                    st.success(f"Extracted {len(text)} characters from {uploaded_file.name}")
                except Exception as e:
                    st.error(f"Failed to extract text: {e}")
    with col2:
        if st.button("Analyze Text") and text_input.strip():
            st.session_state['original_data'] = text_input.strip()
            st.success("Text input saved for analysis.")
    # After input, trigger analysis (simulate analysis for demo)
    if st.session_state['original_data']:
        with st.spinner("Generating medical analysis..."):
            st_lottie(load_lottieurl(LOTTIE_LOADING), height=80, key="main-loading-anim")
            # Simulate analysis (replace with actual call)
            # Here, you would call the orchestrator/agent to generate the report
            # For now, just set a placeholder
            st.session_state['analysis_report'] = "[Analysis will appear here after processing.]"
        st.rerun()

# --- Analysis Report with Download ---
if st.session_state['analysis_report']:
    st.subheader("Medical Analysis Report")
    st.markdown(f"<div style='background:#e3f2fd;padding:1em;border-radius:1em'>{st.session_state['analysis_report']}</div>", unsafe_allow_html=True)
    # --- Download buttons ---
    colpdf, coldocx = st.columns(2)
    with colpdf:
        pdf_buffer = BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=letter)
        textobject = c.beginText(40, 750)
        for line in st.session_state['analysis_report'].splitlines():
            textobject.textLine(line)
        c.drawText(textobject)
        c.showPage()
        c.save()
        pdf_buffer.seek(0)
        st.download_button(
            label="Download as PDF",
            data=pdf_buffer,
            file_name="medical_analysis_report.pdf",
            mime="application/pdf"
        )
    with coldocx:
        docx_buffer = BytesIO()
        doc = Document()
        doc.add_heading('Medical Analysis Report', 0)
        for line in st.session_state['analysis_report'].splitlines():
            doc.add_paragraph(line)
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        st.download_button(
            label="Download as DOCX",
            data=docx_buffer,
            file_name="medical_analysis_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# --- Chat Interface ---
if st.session_state['analysis_report']:
    st.subheader("Chat with the Medical AI")
    st_lottie(load_lottieurl(LOTTIE_CHAT), height=100, key="main-chat-anim")
    st.markdown('<div class="chat-container">', unsafe_allow_html=True)
    for msg in st.session_state['chat_history']:
        if msg['role'] == 'user':
            st.markdown(f'<div class="chat-bubble-user">🧑‍💼 {msg["content"]}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="chat-bubble-ai">🤖 {msg["content"]}</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
    user_question = st.text_input("Type your question about the report, findings, or general medical info:", key="chat-input")
    if st.button("Send", key="send-btn"):
        if user_question.strip():
            st.session_state['chat_history'].append({'role': 'user', 'content': user_question.strip()})
            with st.spinner("AI is thinking..."):
                st_lottie(load_lottieurl(LOTTIE_LOADING), height=80, key="chat-loading")
                result = chat_manager.ask_question(st.session_state['conversation_id'], user_question.strip())
                if result['success']:
                    ai_response = result['response']
                    st.session_state['chat_history'].append({'role': 'assistant', 'content': ai_response})
                else:
                    ai_response = result.get('response', 'Sorry, something went wrong.')
                    st.session_state['chat_history'].append({'role': 'assistant', 'content': ai_response})
            st.rerun()
        else:
            st.warning("Please enter a question.")  